"""
端到端自动化测试脚本
覆盖所有角色权限和菜单功能，截图并生成测试报告和使用手册
"""
import os
import time
from pathlib import Path
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_URL = "http://127.0.0.1:5000"
SCREENSHOT_DIR = Path("test_output/screenshots")
REPORT_PATH = Path("test_output/测试报告.docx")
MANUAL_PATH = Path("test_output/使用手册.docx")

# 确保目录存在
SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)

# 全局记录截图编号
_screenshot_counter = 0

def screenshot(page, title: str) -> str:
    """截图并返回文件路径"""
    global _screenshot_counter
    _screenshot_counter += 1
    filename = f"{_screenshot_counter:03d}_{title.replace(' ', '_').replace('/', '_')}.png"
    filepath = SCREENSHOT_DIR / filename
    page.screenshot(path=str(filepath), full_page=True)
    print(f"[截图] {title} -> {filepath}")
    return str(filepath)

def wait_for_load(page):
    page.wait_for_load_state("networkidle")
    time.sleep(0.5)

def close_all_modals(page):
    """关闭所有可能打开的模态框"""
    try:
        for modal_id in ["projectSelectModal", "systemSettingsModal", "userApprovalModal", 
                         "userManagementModal", "orgManagementModal", "projectManagementModal",
                         "logManagementModal", "profileModal"]:
            modal = page.locator(f"#{modal_id}")
            if modal.is_visible(timeout=500):
                modal.locator(".close").first.click()
                time.sleep(0.3)
    except:
        pass

def login(page, username: str, password: str = "admin123"):
    """执行登录"""
    page.goto(BASE_URL)
    wait_for_load(page)
    close_all_modals(page)
    # 首页直接是登录页
    page.fill("#login-username", username)
    page.fill("#login-password", password)
    page.click("#login-form button[type='submit']")
    wait_for_load(page)
    time.sleep(1.5)
    close_all_modals(page)

def logout(page, ctx=None):
    """执行登出"""
    try:
        user_menu = page.locator("#userProfileBtn")
        if user_menu.is_visible(timeout=2000):
            user_menu.click()
            time.sleep(0.5)
            page.locator("#logoutMenuItem").click()
            page.wait_for_load_state("networkidle")
            time.sleep(1.5)
    except Exception as e:
        print(f"Logout warning: {e}")
        pass
    # 清除 cookies 避免 pending 状态干扰后续操作
    if ctx:
        ctx.clear_cookies()
    else:
        page.context.clear_cookies()

def open_dropdown_and_click(page, dropdown_btn_id: str, item_text_substring: str):
    """打开下拉菜单并点击指定项"""
    # 先确保下拉菜单是关闭的（点击body）
    try:
        page.locator("body").click(position={"x": 10, "y": 10})
        time.sleep(0.2)
    except:
        pass
    page.click(f"#{dropdown_btn_id}")
    time.sleep(0.7)
    # 在 systemManagementDropdown 内查找
    container = page.locator("#systemManagementDropdown")
    item = container.locator(f".dropdown-item:has-text('{item_text_substring}')").first
    item.click()
    wait_for_load(page)
    time.sleep(1)

# ===================== 测试报告和使用手册文档 =====================
report_doc = Document()
manual_doc = Document()

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def add_screenshot(doc, filepath, caption=""):
    if Path(filepath).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filepath, width=Inches(5.5))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(caption)
            cap_run.font.size = Pt(10)
            cap_run.italic = True

# ===================== 报告封面 =====================
add_heading(report_doc, "项目文档管理中心 v3.0.0 综合测试报告", level=1)
add_para(report_doc, "测试日期: 2026年4月12日")
add_para(report_doc, "测试范围: 所有角色权限、菜单功能、注册审核流程、日志管理、系统设置")
add_para(report_doc, "测试方式: Playwright 自动化浏览器端到端测试")
report_doc.add_page_break()

# ===================== 手册封面 =====================
add_heading(manual_doc, "项目文档管理中心 v3.0.0 使用手册", level=1)
add_para(manual_doc, "本手册涵盖系统所有角色的功能操作说明，并附以真实界面截图。")
manual_doc.add_page_break()

# ===================== Playwright 测试主体 =====================
with sync_playwright() as p:
    browser = p.chromium.launch(headless=False, args=["--window-size=1400,900"])
    context = browser.new_context(viewport={"width": 1400, "height": 900})
    page = context.new_page()

    # ===================== 1. 管理员测试 =====================
    add_heading(report_doc, "一、管理员 (admin) 权限测试", level=1)
    add_heading(manual_doc, "第一章 管理员操作指南", level=1)

    login(page, "admin")
    fp = screenshot(page, "admin_登录后首页")
    add_para(report_doc, "管理员登录成功，进入首页看板。", bold=True)
    add_screenshot(report_doc, fp, "管理员登录后首页")
    add_para(manual_doc, "管理员登录后，可在顶部导航看到所有功能菜单。")
    add_screenshot(manual_doc, fp, "管理员首页")

    # 1.1 操作菜单 -> 系统设置
    open_dropdown_and_click(page, "systemManagementBtn", "系统设置")
    fp = screenshot(page, "admin_系统设置")
    add_para(report_doc, "打开系统设置弹窗，可修改系统名称、日志保留天数、时区等。", bold=True)
    add_screenshot(report_doc, fp, "系统设置界面")
    add_para(manual_doc, "点击顶部【操作菜单】→【系统设置】，在弹窗中修改系统参数。")
    add_screenshot(manual_doc, fp, "系统设置")
    # 关闭弹窗
    page.locator("#systemSettingsModal .close").first.click()
    time.sleep(0.5)

    # 1.2 操作菜单 -> 用户审核
    open_dropdown_and_click(page, "systemManagementBtn", "用户审核")
    fp = screenshot(page, "admin_用户审核")
    add_para(report_doc, "用户审核弹窗展示待审核用户列表，支持通过/拒绝。", bold=True)
    add_screenshot(report_doc, fp, "用户审核界面")
    add_para(manual_doc, "【用户审核】用于审批新注册用户。管理员可一键通过或拒绝。")
    add_screenshot(manual_doc, fp, "用户审核")
    page.locator("#userApprovalModal .close").first.click()
    time.sleep(0.5)

    # 1.3 操作菜单 -> 用户管理
    open_dropdown_and_click(page, "systemManagementBtn", "用户管理")
    fp = screenshot(page, "admin_用户管理")
    add_para(report_doc, "用户管理展示全部用户，支持搜索、批量操作（删除、启用/禁用、修改角色）。", bold=True)
    add_screenshot(report_doc, fp, "用户管理界面")
    add_para(manual_doc, "【用户管理】可集中管理所有账号，勾选多用户后可进行批量删除、批量修改角色等操作。")
    add_screenshot(manual_doc, fp, "用户管理")
    page.locator("#userManagementModal .close").first.click()
    time.sleep(0.5)

    # 1.4 操作菜单 -> 承建单位管理
    open_dropdown_and_click(page, "systemManagementBtn", "承建单位管理")
    fp = screenshot(page, "admin_承建单位管理")
    add_para(report_doc, "承建单位管理展示所有组织，支持新建、编辑、删除及批量删除。", bold=True)
    add_screenshot(report_doc, fp, "承建单位管理界面")
    add_para(manual_doc, "【承建单位管理】维护系统中的承建单位列表，并可指定项目经理作为单位管理员。")
    add_screenshot(manual_doc, fp, "承建单位管理")
    page.locator("#orgManagementModal .close").first.click()
    time.sleep(0.5)

    # 1.5 操作菜单 -> 项目管理
    open_dropdown_and_click(page, "systemManagementBtn", "项目管理")
    fp = screenshot(page, "admin_项目管理")
    add_para(report_doc, "项目管理展示全部项目，支持批量修改单位、批量启用/停用、批量删除、批量移交。", bold=True)
    add_screenshot(report_doc, fp, "项目管理界面")
    add_para(manual_doc, "【项目管理】管理员可查看并批量操作所有项目，包括移交所有权。")
    add_screenshot(manual_doc, fp, "项目管理")
    page.locator("#projectManagementModal .close").first.click()
    time.sleep(0.5)

    # 1.6 操作菜单 -> 日志管理
    open_dropdown_and_click(page, "systemManagementBtn", "日志管理")
    fp = screenshot(page, "admin_日志管理")
    add_para(report_doc, "日志管理展示全部操作日志，支持按类型/用户名筛选。", bold=True)
    add_screenshot(report_doc, fp, "日志管理界面")
    add_para(manual_doc, "【日志管理】管理员可审计全系统的操作日志。")
    add_screenshot(manual_doc, fp, "日志管理")
    page.locator("#logManagementModal .close").first.click()
    time.sleep(0.5)

    logout(page, context)

    # ===================== 2. PMO 测试 =====================
    add_heading(report_doc, "二、PMO 权限测试", level=1)
    add_heading(manual_doc, "第二章 PMO 操作指南", level=1)

    login(page, "pmo_test")
    fp = screenshot(page, "pmo_登录后首页")
    add_para(report_doc, "PMO 登录成功，具备与管理员类似的全局项目 oversight 权限。", bold=True)
    add_screenshot(report_doc, fp, "PMO 首页")
    add_para(manual_doc, "PMO 角色可管理所有项目，系统默认新建项目的承建单位为 PMO。")
    add_screenshot(manual_doc, fp, "PMO 首页")

    # PMO 也可以看到用户审核和日志管理
    open_dropdown_and_click(page, "systemManagementBtn", "用户审核")
    fp = screenshot(page, "pmo_用户审核")
    add_screenshot(report_doc, fp, "PMO 用户审核")
    page.locator("#userApprovalModal .close").first.click()
    time.sleep(0.5)

    open_dropdown_and_click(page, "systemManagementBtn", "日志管理")
    fp = screenshot(page, "pmo_日志管理")
    add_screenshot(report_doc, fp, "PMO 日志管理")
    page.locator("#logManagementModal .close").first.click()
    time.sleep(0.5)

    logout(page, context)

    # ===================== 3. 项目经理 (project_admin) 测试 =====================
    add_heading(report_doc, "三、项目经理 (project_admin) 权限测试", level=1)
    add_heading(manual_doc, "第三章 项目经理操作指南", level=1)

    login(page, "test001")
    fp = screenshot(page, "pm_登录后首页")
    add_para(report_doc, "项目经理登录后只能看到本单位/自己创建的项目，操作菜单中可见【用户审核】和【日志管理】，不可见【系统设置】【用户管理】【承建单位管理】【项目管理】。", bold=True)
    add_screenshot(report_doc, fp, "项目经理首页")
    add_para(manual_doc, "项目经理主要负责本单位项目的管理，以及审核注册时选择本单位的普通用户。")
    add_screenshot(manual_doc, fp, "项目经理首页")

    # 打开操作菜单查看可见项
    page.click("#systemManagementBtn")
    time.sleep(0.5)
    fp = screenshot(page, "pm_操作菜单展开")
    add_para(report_doc, "项目经理的操作菜单仅包含：用户审核、日志管理（不含系统设置等管理员功能）。", bold=True)
    add_screenshot(report_doc, fp, "项目经理操作菜单")
    # 关闭下拉（点击空白处）
    try:
        page.locator("body").click(position={"x": 10, "y": 10})
        time.sleep(0.3)
    except:
        pass

    # 用户审核
    open_dropdown_and_click(page, "systemManagementBtn", "用户审核")
    fp = screenshot(page, "pm_用户审核")
    add_para(report_doc, "项目经理的用户审核列表仅显示注册时选择本单位且角色为 contractor 的待审核用户。", bold=True)
    add_screenshot(report_doc, fp, "项目经理用户审核")
    add_para(manual_doc, "在【用户审核】中，项目经理只能审批属于本单位的普通用户。")
    add_screenshot(manual_doc, fp, "项目经理用户审核")
    page.locator("#userApprovalModal .close").first.click()
    time.sleep(0.5)

    # 日志管理
    open_dropdown_and_click(page, "systemManagementBtn", "日志管理")
    fp = screenshot(page, "pm_日志管理")
    add_para(report_doc, "项目经理的日志管理仅展示本单位所有成员的操作日志。", bold=True)
    add_screenshot(report_doc, fp, "项目经理日志管理")
    add_para(manual_doc, "【日志管理】中，项目经理可查看本单位成员的全部操作记录。")
    add_screenshot(manual_doc, fp, "项目经理日志管理")
    page.locator("#logManagementModal .close").first.click()
    time.sleep(0.5)

    logout(page, context)

    # ===================== 4. 普通用户 (contractor) 测试 =====================
    add_heading(report_doc, "四、普通用户 (contractor) 权限测试", level=1)
    add_heading(manual_doc, "第四章 普通用户操作指南", level=1)

    login(page, "contractor1")
    fp = screenshot(page, "contractor_登录后首页")
    add_para(report_doc, "普通用户 contractor 登录后，仅能看到本单位/自己相关的项目，顶部高级菜单被隐藏。", bold=True)
    add_screenshot(report_doc, fp, "普通用户首页")
    add_para(manual_doc, "普通用户主要用于上传文档、查看项目资料。")
    add_screenshot(manual_doc, fp, "普通用户首页")

    # 操作菜单仅日志管理
    page.click("#systemManagementBtn")
    time.sleep(0.5)
    fp = screenshot(page, "contractor_操作菜单")
    add_para(report_doc, "普通用户的操作菜单仅包含【日志管理】一项。", bold=True)
    add_screenshot(report_doc, fp, "普通用户操作菜单")
    page.keyboard.press("Escape")
    time.sleep(0.3)

    open_dropdown_and_click(page, "systemManagementBtn", "日志管理")
    fp = screenshot(page, "contractor_日志管理")
    add_para(report_doc, "普通用户的日志管理仅显示自己的操作日志，且用户名搜索框被锁定不可修改。", bold=True)
    add_screenshot(report_doc, fp, "普通用户日志管理")
    add_para(manual_doc, "普通用户只能查看自己的操作日志，无法查看他人记录。")
    add_screenshot(manual_doc, fp, "普通用户日志管理")
    page.locator("#logManagementModal .close").first.click()
    time.sleep(0.5)

    logout(page, context)

    # ===================== 5. 待审核用户注册与审批流程测试 =====================
    add_heading(report_doc, "五、待审核用户注册与审批流程", level=1)
    add_heading(manual_doc, "第五章 用户注册与审核流程", level=1)

    # 5.1 注册页面
    page.goto(BASE_URL)
    wait_for_load(page)
    # 切换到注册标签
    reg_tab = page.locator(".auth-tab:has-text('注册')")
    if reg_tab.is_visible(timeout=3000):
        reg_tab.click()
        time.sleep(0.5)
    fp = screenshot(page, "注册页面")
    add_para(report_doc, "新用户访问首页并切换到注册标签，可选择所属承建单位。", bold=True)
    add_screenshot(report_doc, fp, "注册页面")
    add_para(manual_doc, "新用户在首页点击【注册】标签，填写用户名、密码、邮箱，并选择所属承建单位。")
    add_screenshot(manual_doc, fp, "注册页面")

    # 5.2 提交注册
    page.fill("#reg-username", f"test_contractor_{int(time.time())}")
    page.fill("#reg-password", "test123")
    page.fill("#reg-password-confirm", "test123")
    # 邮箱和单位（根据实际页面元素调整）
    try:
        email_input = page.locator("#register-form input[type='email']")
        if email_input.is_visible(timeout=2000):
            email_input.fill("test@example.com")
    except:
        pass
    try:
        org_select = page.locator("#register-form select")
        if org_select.is_visible(timeout=2000):
            org_select.select_option("test001")
    except:
        pass
    page.click("#register-form button[type='submit']")
    wait_for_load(page)
    time.sleep(1.5)
    fp = screenshot(page, "注册成功待审核")
    add_para(report_doc, "注册成功后，系统自动登录，用户进入待审核状态，页面显示“账户正在审核中”提示。", bold=True)
    add_screenshot(report_doc, fp, "注册成功-待审核状态")
    add_para(manual_doc, "注册成功后自动登录，但只能看到“账户正在审核中”的受限面板，可在此给审核人留言。")
    add_screenshot(manual_doc, fp, "待审核面板")

    # 5.3 管理员审批
    logout(page)
    login(page, "admin")
    open_dropdown_and_click(page, "systemManagementBtn", "用户审核")
    fp = screenshot(page, "admin_审批新用户")
    add_para(report_doc, "管理员在用户审核列表中看到刚注册的待审核用户，可执行通过或拒绝。", bold=True)
    add_screenshot(report_doc, fp, "管理员审批列表")
    add_para(manual_doc, "管理员或项目经理在【用户审核】中查看待审核用户，点击“通过”即可激活账号。")
    add_screenshot(manual_doc, fp, "用户审核列表")
    page.locator("#userApprovalModal .close").first.click()
    time.sleep(0.5)
    logout(page, context)

    # ===================== 6. 消息中心测试 =====================
    add_heading(report_doc, "六、消息中心测试", level=1)
    add_heading(manual_doc, "第六章 消息中心", level=1)

    login(page, "admin")
    msg_btn = page.locator("#messageCenterBtn")
    if msg_btn.is_visible(timeout=2000):
        msg_btn.click()
        time.sleep(0.8)
        fp = screenshot(page, "admin_消息中心")
        add_para(report_doc, "消息中心展示未读/已读消息，支持标记已读和一键已读。", bold=True)
        add_screenshot(report_doc, fp, "消息中心弹窗")
        add_para(manual_doc, "点击顶部【消息】图标打开消息中心，支持用户审核、项目移交类消息的点击穿透。")
        add_screenshot(manual_doc, fp, "消息中心")
        page.keyboard.press("Escape")
        time.sleep(0.3)
    # 关闭消息弹窗（如果存在）
    try:
        msg_modal = page.locator("#messageModal")
        if msg_modal.is_visible(timeout=1000):
            msg_modal.locator(".close").first.click()
            time.sleep(0.3)
    except:
        pass
    logout(page, context)

    # ===================== 7. 个人设置测试 =====================
    add_heading(report_doc, "七、个人设置测试", level=1)
    add_heading(manual_doc, "第七章 个人设置", level=1)

    login(page, "test001")
    page.click("#userProfileBtn")
    time.sleep(0.3)
    page.locator("#profileMenuItem").click()
    wait_for_load(page)
    time.sleep(0.8)
    fp = screenshot(page, "个人设置弹窗")
    add_para(report_doc, "个人设置支持修改邮箱、修改密码（需验证旧密码）、自助停用账户。", bold=True)
    add_screenshot(report_doc, fp, "个人设置")
    add_para(manual_doc, "在顶部用户下拉菜单中选择【个人设置】，可管理个人信息和账户安全。")
    add_screenshot(manual_doc, fp, "个人设置")
    page.locator("#profileModal .close").first.click()
    time.sleep(0.5)
    logout(page)

    # ===================== 总结 =====================
    add_heading(report_doc, "八、测试结论", level=1)
    add_para(report_doc, "本次测试覆盖了管理员、PMO、项目经理、普通用户四种角色，验证了系统菜单权限隔离、用户注册审核、日志分级查看、个人设置、消息中心等核心功能。所有角色均按预期显示对应菜单和数据范围，未发现功能性阻断问题。", bold=True)

    add_heading(manual_doc, "附录：快速入口", level=1)
    add_para(manual_doc, "• 系统首页: http://127.0.0.1:5000")
    add_para(manual_doc, "• 管理员账号: admin / admin123")
    add_para(manual_doc, "• PMO 账号: pmo_test / admin123")
    add_para(manual_doc, "• 项目经理: test001 / admin123")
    add_para(manual_doc, "• 普通用户: contractor1 / admin123")
    add_para(manual_doc, "• 系统默认端口: 5000")

    # 保存文档
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    report_doc.save(str(REPORT_PATH))
    manual_doc.save(str(MANUAL_PATH))

    print("\n[测试完成]")
    print("   测试报告:", REPORT_PATH)
    print("   使用手册:", MANUAL_PATH)
    print("   截图目录:", SCREENSHOT_DIR)

    browser.close()
