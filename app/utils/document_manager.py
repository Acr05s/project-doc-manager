"""项目文档管理器 - 模块化版本

本模块组合所有子模块，提供统一的文档管理接口。
每个子模块独立运行，某个模块加载错误不会影响其他模块。

模块结构：
- base: 基础配置和工具函数
- cache_manager: 缓存管理
- operation_logger: 操作日志
- doc_naming: 文档命名和编号
- folder_manager: 文件夹管理
- document_list: 文档清单管理
- archive_manager: 归档管理
- export_manager: 导出管理
- requirements_loader: 需求加载
- project_manager: 项目管理
- image_analyzer: 图像分析
- document_uploader: 文档上传
"""

import logging
from typing import Dict, Any, Optional, List

# 导入各模块
from .base import DocumentConfig, get_config, setup_logging, get_base_dir

logger = setup_logging(__name__)

# 尝试导入各模块，如果失败则记录警告但继续运行
try:
    from .cache_manager import CacheManager
    _cache_manager_available = True
except Exception as e:
    logger.warning(f"缓存管理模块加载失败: {e}")
    _cache_manager_available = False

try:
    from .operation_logger import OperationLogger
    _operation_logger_available = True
except Exception as e:
    logger.warning(f"操作日志模块加载失败: {e}")
    _operation_logger_available = False

try:
    from .doc_naming import DocumentNamer
    _doc_naming_available = True
except Exception as e:
    logger.warning(f"文档命名模块加载失败: {e}")
    _doc_naming_available = False

try:
    from .folder_manager import FolderManager
    _folder_manager_available = True
except Exception as e:
    logger.warning(f"文件夹管理模块加载失败: {e}")
    _folder_manager_available = False

try:
    from .document_list import DocumentListManager
    _document_list_available = True
except Exception as e:
    logger.warning(f"文档清单模块加载失败: {e}")
    _document_list_available = False

try:
    from .archive_manager import ArchiveManager
    _archive_manager_available = True
except Exception as e:
    logger.warning(f"归档管理模块加载失败: {e}")
    _archive_manager_available = False

try:
    from .export_manager import ExportManager
    _export_manager_available = True
except Exception as e:
    logger.warning(f"导出管理模块加载失败: {e}")
    _export_manager_available = False

try:
    from .requirements_loader import RequirementsLoader
    _requirements_loader_available = True
except Exception as e:
    logger.warning(f"需求加载模块加载失败: {e}")
    _requirements_loader_available = False

try:
    from .project_manager import ProjectManager
    _project_manager_available = True
except Exception as e:
    logger.warning(f"项目管理模块加载失败: {e}")
    _project_manager_available = False

try:
    from .image_analyzer import ImageAnalyzer
    _image_analyzer_available = True
except Exception as e:
    logger.warning(f"图像分析模块加载失败: {e}")
    _image_analyzer_available = False

try:
    from .document_uploader import DocumentUploader
    _document_uploader_available = True
except Exception as e:
    logger.warning(f"文档上传模块加载失败: {e}")
    _document_uploader_available = False

try:
    from .project_data_manager import ProjectDataManager
    _project_data_manager_available = True
except Exception as e:
    logger.warning(f"项目数据管理模块加载失败: {e}")
    _project_data_manager_available = False


class DocumentManager:
    """项目文档管理器 - 主类
    
    组合所有子模块，提供统一的文档管理接口。
    每个子模块通过属性访问，如果模块加载失败则返回None。
    """
    
    def __init__(self, config: Optional[Dict] = None):
        """初始化文档管理器
        
        Args:
            config: 配置字典
        """
        # 初始化配置
        self.config = get_config(config)
        self.name = "项目文档管理中心"
        self.version = "2.1.0"  # 模块化版本
        
        # 文档数据库（内存缓存）
        self.documents_db = {}
        
        # 待确认文件
        self.pending_files = {}
        
        # 已确认文件
        self.confirmed_files = {}
        
        # 上传会话
        self.upload_sessions = {}
        
        # 初始化各子模块
        self._init_modules()
        
        logger.info(f"文档管理器初始化完成 (v{self.version})")
    
    def _init_modules(self):
        """初始化所有子模块"""
        # 缓存管理
        if _cache_manager_available:
            try:
                self.cache = CacheManager(self.config)
            except Exception as e:
                logger.error(f"初始化缓存管理模块失败: {e}")
                self.cache = None
        else:
            self.cache = None
        
        # 操作日志
        if _operation_logger_available:
            try:
                self.logger = OperationLogger(self.config)
            except Exception as e:
                logger.error(f"初始化操作日志模块失败: {e}")
                self.logger = None
        else:
            self.logger = None
        
        # 文件夹管理
        if _folder_manager_available:
            try:
                self.folders = FolderManager(self.config)
            except Exception as e:
                logger.error(f"初始化文件夹管理模块失败: {e}")
                self.folders = None
        else:
            self.folders = None
        
        # 文档命名
        if _doc_naming_available:
            try:
                self.naming = DocumentNamer(self.config)
            except Exception as e:
                logger.error(f"初始化文档命名模块失败: {e}")
                self.naming = None
        else:
            self.naming = None
        
        # 文档清单
        if _document_list_available and self.folders:
            try:
                self.doc_lists = DocumentListManager(self.config, self.folders)
            except Exception as e:
                logger.error(f"初始化文档清单模块失败: {e}")
                self.doc_lists = None
        else:
            self.doc_lists = None
        
        # 归档管理
        if _archive_manager_available and self.folders and self.doc_lists:
            try:
                self.archive = ArchiveManager(self.config, self.folders, self.doc_lists)
            except Exception as e:
                logger.error(f"初始化归档管理模块失败: {e}")
                self.archive = None
        else:
            self.archive = None
        
        # 导出管理
        if _export_manager_available and self.folders and self.doc_lists:
            try:
                self.exporter = ExportManager(self.config, self.folders, self.doc_lists)
            except Exception as e:
                logger.error(f"初始化导出管理模块失败: {e}")
                self.exporter = None
        else:
            self.exporter = None
        
        # 需求加载
        if _requirements_loader_available:
            try:
                self.requirements = RequirementsLoader(self.config)
            except Exception as e:
                logger.error(f"初始化需求加载模块失败: {e}")
                self.requirements = None
        else:
            self.requirements = None
        
        # 项目管理
        if _project_manager_available and self.folders:
            try:
                self.projects = ProjectManager(self.config, self.folders)
            except Exception as e:
                logger.error(f"初始化项目管理模块失败: {e}")
                self.projects = None
        else:
            self.projects = None
        
        # 图像分析
        if _image_analyzer_available:
            try:
                self.analyzer = ImageAnalyzer(self.config)
            except Exception as e:
                logger.error(f"初始化图像分析模块失败: {e}")
                self.analyzer = None
        else:
            self.analyzer = None
        
        # 文档上传
        if _document_uploader_available and self.folders:
            try:
                self.uploader = DocumentUploader(self.config, self.folders)
            except Exception as e:
                logger.error(f"初始化文档上传模块失败: {e}")
                self.uploader = None
        else:
            self.uploader = None
        
        # 项目数据管理
        if _project_data_manager_available:
            try:
                self.data_manager = ProjectDataManager(self.config)
            except Exception as e:
                logger.error(f"初始化项目数据管理模块失败: {e}")
                self.data_manager = None
        else:
            self.data_manager = None
    
    def get_status(self) -> Dict[str, Any]:
        """获取模块状态
        
        Returns:
            Dict: 各模块的加载状态
        """
        return {
            'version': self.version,
            'modules': {
                'cache': _cache_manager_available and self.cache is not None,
                'logger': _operation_logger_available and self.logger is not None,
                'folders': _folder_manager_available and self.folders is not None,
                'naming': _doc_naming_available and self.naming is not None,
                'doc_lists': _document_list_available and self.doc_lists is not None,
                'archive': _archive_manager_available and self.archive is not None,
                'exporter': _export_manager_available and self.exporter is not None,
                'requirements': _requirements_loader_available and self.requirements is not None,
                'projects': _project_manager_available and self.projects is not None,
                'analyzer': _image_analyzer_available and self.analyzer is not None,
                'uploader': _document_uploader_available and self.uploader is not None,
            }
        }
    
    @property
    def projects_base_folder(self):
        """获取项目基础文件夹
        
        Returns:
            Path: 项目基础文件夹路径
        """
        return self.config.projects_base_folder
    
    # ==================== 便捷方法 ====================
    # 下面提供一些常用方法的便捷访问
    
    def log_operation(self, operation: str, details: str = '', 
                     status: str = 'success', project: str = None):
        """记录操作日志"""
        if self.logger:
            self.logger.log(operation, details, status, project)
    
    def get_projects_list(self) -> List[Dict]:
        """获取项目列表"""
        if self.projects:
            return self.projects.list_all()
        return []
    
    def create_project(self, name: str, description: str = '',
                      requirements_file: Optional[str] = None,
                      party_a: str = '', party_b: str = '',
                      supervisor: str = '', manager: str = '',
                      duration: str = '') -> Dict:
        """创建项目"""
        if self.projects:
            return self.projects.create(name, description, requirements_file,
                                       party_a, party_b, supervisor, manager, duration)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def import_project_json(self, project_config: Dict, new_name: str = None) -> Dict:
        """从JSON配置导入项目
        
        Args:
            project_config: 项目配置JSON
            new_name: 可选的新项目名称（如果不指定则使用配置中的名称）
            
        Returns:
            Dict: 导入结果
        """
        try:
            if not self.projects:
                return {'status': 'error', 'message': '项目管理模块不可用'}
            
            # 获取项目名称
            project_name = new_name or project_config.get('name', '导入项目')
            
            # 生成新的项目ID
            from datetime import datetime
            new_project_id = f"project_{datetime.now().strftime('%Y%m%d%H%M%S')}"
            
            # 检查项目名称是否已存在
            existing_projects = self.projects.list_all()
            for proj in existing_projects:
                if proj.get('name') == project_name:
                    # 如果项目名已存在，添加时间戳后缀
                    project_name = f"{project_name}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                    break
            
            # 更新配置
            project_config['id'] = new_project_id
            project_config['name'] = project_name
            project_config['created_time'] = datetime.now().isoformat()
            project_config['updated_time'] = datetime.now().isoformat()
            
            # 创建项目目录结构
            self.data_manager.create_project_structure(project_name)
            
            # 保存项目配置
            self.data_manager.save_full_config(project_name, project_config)
            
            # 添加到项目索引
            self.projects.add_to_index(
                project_id=new_project_id,
                name=project_name,
                description=project_config.get('description', ''),
                created_time=project_config['created_time']
            )
            
            logger.info(f"项目导入成功: {project_name} (ID: {new_project_id})")
            
            return {
                'status': 'success',
                'project_id': new_project_id,
                'project_name': project_name,
                'message': '项目导入成功'
            }
            
        except Exception as e:
            logger.error(f"导入项目失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return {'status': 'error', 'message': str(e)}
    
    def load_project(self, project_id: str) -> Dict:
        """加载项目"""
        if self.projects:
            project = self.projects.load(project_id)
            if project:
                return {'status': 'success', 'project': project}
            return {'status': 'error', 'message': '项目不存在'}
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def save_project(self, project_config: Dict) -> Dict:
        """保存项目"""
        if self.projects:
            project_id = project_config.get('id')
            return self.projects.save(project_id, project_config)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def delete_project(self, project_id: str, permanent: bool = False) -> Dict:
        """删除项目（软删除）"""
        if self.projects:
            return self.projects.delete(project_id, permanent)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def restore_project(self, project_id: str) -> Dict:
        """恢复已删除的项目"""
        if self.projects:
            return self.projects.restore(project_id)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def get_deleted_projects(self) -> List[Dict]:
        """获取已删除项目列表"""
        if self.projects:
            return self.projects.list_deleted()
        return []
    
    def load_requirements(self, file_path: str) -> Dict:
        """加载需求"""
        if self.requirements:
            return self.requirements.load(file_path)
        return {'cycles': [], 'documents': {}}

    def save_requirements_config(self, config: Dict, source_filename: str = None) -> Dict:
        """保存文档需求配置为独立文件

        Args:
            config: 文档需求配置（包含cycles和documents）
            source_filename: 源文件名（用于生成配置名称）

        Returns:
            Dict: 保存结果，包含requirements_id
        """
        logger.info(f"[DEBUG] save_requirements_config 被调用, source_filename: {source_filename}")
        try:
            from datetime import datetime
            import json
            from pathlib import Path

            # 生成配置ID和文件名
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            requirements_id = f"req_{timestamp}"
            logger.info(f"[DEBUG] 生成的 requirements_id: {requirements_id}")

            # 使用源文件名作为配置名称（去掉扩展名）
            if source_filename:
                config_name = Path(source_filename).stem
            else:
                config_name = f"需求配置_{timestamp}"

            # 确保requirements目录存在
            requirements_dir = self.config.projects_folder / 'requirements'
            requirements_dir.mkdir(parents=True, exist_ok=True)

            # 保存配置
            config_file = requirements_dir / f"{requirements_id}.json"
            config_data = {
                'id': requirements_id,
                'name': config_name,
                'created_time': datetime.now().isoformat(),
                'cycles': config.get('cycles', []),
                'documents': config.get('documents', {})
            }

            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(config_data, f, ensure_ascii=False, indent=2)

            logger.info(f"文档需求配置已保存: {config_file}")
            
            result = {
                'status': 'success',
                'requirements_id': requirements_id,
                'name': config_name
            }
            logger.info(f"[DEBUG] 返回结果: {result}")
            return result

        except Exception as e:
            logger.error(f"保存文档需求配置失败: {e}")
            return {'status': 'error', 'message': str(e)}

    def apply_requirements_to_project(self, project_id: str, requirements_id: str) -> Dict:
        """将文档需求配置应用到项目

        Args:
            project_id: 项目ID
            requirements_id: 需求配置ID

        Returns:
            Dict: 应用结果
        """
        try:
            import json
            from pathlib import Path
            from datetime import datetime

            logger.info(f"[DEBUG] apply_requirements_to_project 被调用: project_id={project_id}, requirements_id={requirements_id}")

            # 读取需求配置文件
            requirements_dir = self.config.projects_folder / 'requirements'
            req_file = requirements_dir / f"{requirements_id}.json"
            
            logger.info(f"[DEBUG] 查找需求配置文件: {req_file}")
            logger.info(f"[DEBUG] 文件是否存在: {req_file.exists()}")

            if not req_file.exists():
                return {'status': 'error', 'message': f'需求配置不存在: {req_file}'}

            with open(req_file, 'r', encoding='utf-8') as f:
                req_config = json.load(f)
            
            logger.info(f"[DEBUG] 需求配置已加载: cycles={len(req_config.get('cycles', []))}")

            # 获取项目信息
            project = self.projects.load(project_id) if self.projects else None
            if not project:
                return {'status': 'error', 'message': '项目不存在'}
            
            logger.info(f"[DEBUG] 项目已加载: {project}")

            project_name = project.get('name', project_id)
            logger.info(f"[DEBUG] 项目名称: {project_name}")

            # 确保项目目录存在
            project_folder = self.config.projects_folder / project_name
            project_folder.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"[DEBUG] 项目目录: {project_folder}, 是否存在: {project_folder.exists()}")

            # 加载旧的requirements.json（如果存在）
            old_req_file = project_folder / 'requirements.json'
            old_required_docs = {}
            if old_req_file.exists():
                try:
                    with open(old_req_file, 'r', encoding='utf-8') as f:
                        old_config = json.load(f)
                    old_documents = old_config.get('documents', {})
                    for cycle, cycle_data in old_documents.items():
                        old_required_docs[cycle] = cycle_data.get('required_docs', [])
                except Exception as e:
                    logger.warning(f"[DEBUG] 读取旧配置失败: {e}")

            # 构建文档名称映射（旧名称 -> 新名称）
            doc_name_mapping = {}
            new_documents = req_config.get('documents', {})
            for cycle, new_cycle_data in new_documents.items():
                if cycle not in old_required_docs:
                    continue
                
                old_docs = old_required_docs[cycle]
                new_docs = new_cycle_data.get('required_docs', [])
                
                # 对比新旧文档列表，找出名称变更
                for i, old_doc in enumerate(old_docs):
                    if i < len(new_docs):
                        old_name = old_doc.get('name') if isinstance(old_doc, dict) else str(old_doc)
                        new_name = new_docs[i].get('name') if isinstance(new_docs[i], dict) else str(new_docs[i])
                        if old_name and new_name and old_name != new_name:
                            doc_name_mapping[(project_name, cycle, old_name)] = new_name

            logger.info(f"[DEBUG] 文档名称映射: {doc_name_mapping}")

            # 同步更新 documents_index.json 中的 doc_name
            if doc_name_mapping:
                data_dir = project_folder / 'data'
                index_file = data_dir / 'documents_index.json'
                
                if index_file.exists():
                    try:
                        with open(index_file, 'r', encoding='utf-8') as f:
                            doc_index = json.load(f)
                        
                        documents = doc_index.get('documents', {})
                        updated_count = 0
                        
                        for doc_id, doc_info in documents.items():
                            cycle = doc_info.get('cycle', '')
                            doc_name = doc_info.get('doc_name', '')
                            mapping_key = (project_name, cycle, doc_name)
                            
                            if mapping_key in doc_name_mapping:
                                old_name = doc_info['doc_name']
                                doc_info['doc_name'] = doc_name_mapping[mapping_key]
                                updated_count += 1
                                logger.info(f"[DEBUG] 更新文档名称: {old_name} -> {doc_info['doc_name']}")
                        
                        if updated_count > 0:
                            # 备份原文件
                            backup_file = data_dir / 'documents_index_backup.json'
                            with open(backup_file, 'w', encoding='utf-8') as f:
                                json.dump(doc_index, f, ensure_ascii=False, indent=2)
                            
                            # 保存更新后的文件
                            with open(index_file, 'w', encoding='utf-8') as f:
                                json.dump(doc_index, f, ensure_ascii=False, indent=2)
                            
                            logger.info(f"[DEBUG] 已更新 {updated_count} 个文档名称")
                    except Exception as e:
                        logger.error(f"[DEBUG] 更新文档索引失败: {e}")
                        import traceback
                        logger.error(f"[DEBUG] 错误堆栈: {traceback.format_exc()}")

            # 将需求配置复制到项目目录
            target_file = project_folder / 'requirements.json'
            config_to_save = {
                'requirements_id': requirements_id,
                'name': req_config.get('name', ''),
                'applied_time': datetime.now().isoformat(),
                'cycles': req_config.get('cycles', []),
                'documents': req_config.get('documents', {})
            }

            with open(target_file, 'w', encoding='utf-8') as f:
                json.dump(config_to_save, f, ensure_ascii=False, indent=2)
            
            logger.info(f"[DEBUG] 配置文件已保存到: {target_file}")
            logger.info(f"[DEBUG] 文件是否写入成功: {target_file.exists()}")

            logger.info(f"需求配置已应用到项目 {project_id}: {target_file}")

            return {
                'status': 'success',
                'message': '文档需求配置已应用到项目',
                'cycles': req_config.get('cycles', []),
                'documents': req_config.get('documents', {})
            }

        except Exception as e:
            logger.error(f"[DEBUG] 应用需求配置失败: {e}")
            import traceback
            logger.error(f"[DEBUG] 错误堆栈: {traceback.format_exc()}")
            return {'status': 'error', 'message': str(e)}

    def list_requirements_configs(self) -> Dict:
        """获取所有文档需求配置列表

        Returns:
            Dict: 配置列表
        """
        try:
            import json
            from pathlib import Path

            requirements_dir = self.config.projects_folder / 'requirements'
            requirements_dir.mkdir(parents=True, exist_ok=True)

            configs = []
            for config_file in requirements_dir.glob('*.json'):
                try:
                    with open(config_file, 'r', encoding='utf-8') as f:
                        config = json.load(f)
                        configs.append({
                            'id': config.get('id', config_file.stem),
                            'name': config.get('name', config_file.stem),
                            'created_time': config.get('created_time', ''),
                            'cycle_count': len(config.get('cycles', [])),
                            'doc_count': sum(len(docs.get('required_docs', []))
                                           for docs in config.get('documents', {}).values())
                        })
                except Exception as e:
                    logger.warning(f"读取配置文件失败 {config_file}: {e}")

            # 按创建时间倒序
            configs.sort(key=lambda x: x.get('created_time', ''), reverse=True)

            return {
                'status': 'success',
                'configs': configs
            }

        except Exception as e:
            logger.error(f"获取需求配置列表失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def load_document_list(self, project_name: str) -> Optional[Dict]:
        """加载文档清单"""
        if self.doc_lists:
            # 先尝试加载文档清单文件
            doc_list = self.doc_lists.load(project_name)
            if doc_list:
                return doc_list
        
        # 如果文档清单文件不存在，尝试从项目数据管理器中加载数据
        if self.data_manager:
            config = self.data_manager.load_full_config(project_name)
            if config:
                # 转换为文档清单格式
                doc_list = {
                    'version': '2.0',
                    'project_name': project_name,
                    'created_time': config.get('created_time', self._get_timestamp()),
                    'updated_time': config.get('updated_time', self._get_timestamp()),
                    'project_info': {
                        'name': config.get('name', project_name),
                        'description': config.get('description', ''),
                        'party_a': config.get('party_a', ''),
                        'party_b': config.get('party_b', ''),
                        'supervisor': config.get('supervisor', ''),
                        'manager': config.get('manager', ''),
                        'duration': config.get('duration', '')
                    },
                    'cycles': []
                }
                
                # 添加周期和文档
                cycles = config.get('cycles', [])
                documents = config.get('documents', {})
                
                for cycle in cycles:
                    cycle_entry = {
                        'name': cycle,
                        'description': '',
                        'documents': []
                    }
                    
                    # 添加该周期的文档
                    cycle_docs = documents.get(cycle, {})
                    required_docs = cycle_docs.get('required_docs', [])
                    uploaded_docs = cycle_docs.get('uploaded_docs', [])
                    
                    # 创建文档映射
                    doc_map = {}
                    for req_doc in required_docs:
                        doc_map[req_doc.get('name')] = req_doc
                    
                    # 添加文档
                    for req_doc in required_docs:
                        doc_name = req_doc.get('name')
                        doc_entry = {
                            'name': doc_name,
                            'requirement': req_doc.get('requirement', ''),
                            'files': [],
                            'status': req_doc.get('status', 'pending')
                        }
                        
                        # 添加上传的文件
                        for uploaded_doc in uploaded_docs:
                            if uploaded_doc.get('doc_name') == doc_name:
                                file_info = {
                                    'filename': uploaded_doc.get('filename'),
                                    'original_filename': uploaded_doc.get('original_filename'),
                                    'file_path': uploaded_doc.get('file_path'),
                                    'upload_time': uploaded_doc.get('upload_time'),
                                    'source': uploaded_doc.get('source'),
                                    'file_size': uploaded_doc.get('file_size')
                                }
                                doc_entry['files'].append(file_info)
                        
                        # 更新状态
                        if doc_entry['files']:
                            doc_entry['status'] = 'completed' if len(doc_entry['files']) > 0 else 'pending'
                        
                        cycle_entry['documents'].append(doc_entry)
                    
                    # 添加未在required_docs中但已上传的文档
                    for uploaded_doc in uploaded_docs:
                        doc_name = uploaded_doc.get('doc_name')
                        if doc_name and not any(d.get('name') == doc_name for d in cycle_entry['documents']):
                            doc_entry = {
                                'name': doc_name,
                                'requirement': '',
                                'files': [{
                                    'filename': uploaded_doc.get('filename'),
                                    'original_filename': uploaded_doc.get('original_filename'),
                                    'file_path': uploaded_doc.get('file_path'),
                                    'upload_time': uploaded_doc.get('upload_time'),
                                    'source': uploaded_doc.get('source'),
                                    'file_size': uploaded_doc.get('file_size')
                                }],
                                'status': 'completed'
                            }
                            cycle_entry['documents'].append(doc_entry)
                    
                    doc_list['cycles'].append(cycle_entry)
                
                logger.info(f"从项目数据管理器加载文档清单: {project_name}")
                return doc_list
        
        return None
    
    def _get_timestamp(self) -> str:
        """获取当前时间戳"""
        from datetime import datetime
        return datetime.now().isoformat()
    
    def save_document_list(self, project_name: str, doc_list: Dict) -> Dict:
        """保存文档清单"""
        if self.doc_lists:
            return self.doc_lists.save(project_name, doc_list)
        return {'status': 'error', 'message': '文档清单模块不可用'}
    
    def export_documents_package(self, project_name: str) -> Dict:
        """导出文档包"""
        if self.exporter:
            return self.exporter.export_documents_package(project_name)
        return {'status': 'error', 'message': '导出模块不可用'}
    
    def archive_document(self, project_name: str, cycle_name: str, doc_name: str,
                        source_file, original_filename: str) -> Dict:
        """归档文档"""
        if self.archive:
            return self.archive.archive(project_name, cycle_name, doc_name, 
                                       source_file, original_filename)
        return {'status': 'error', 'message': '归档模块不可用'}
    
    def upload_document(self, file, cycle: str, doc_name: str, **kwargs) -> Dict:
        """上传文档
        
        Args:
            file: 文件对象
            cycle: 项目周期
            doc_name: 文档名称
            **kwargs: 其他参数，包括 project_name
            
        Returns:
            Dict: 上传结果
        """
        if self.uploader:
            return self.uploader.upload(file, cycle, doc_name, **kwargs)
        return {'status': 'error', 'message': '上传模块不可用'}
    
    def detect_signature(self, image_path: str) -> tuple:
        """检测签字"""
        if self.analyzer:
            return self.analyzer.detect_signature(image_path)
        return False, 0.0
    
    def detect_seal(self, image_path: str) -> tuple:
        """检测盖章"""
        if self.analyzer:
            return self.analyzer.detect_seal(image_path)
        return False, 0.0
    
    # ==================== 兼容旧版API ====================
    # 以下方法保留用于兼容旧版本代码
    
    def remove_leading_number(self, filename: str) -> str:
        """删除文件名中第一个中文字符前的内容"""
        if self.naming:
            return self.naming.remove_leading_number(filename)
        return filename
    
    def generate_doc_number(self, project_name: str, cycle_index: int, 
                           doc_index: int, doc_type: str = None) -> str:
        """生成文档编号"""
        if self.naming:
            return self.naming.generate_doc_number(project_name, cycle_index, 
                                                   doc_index, doc_type)
        return f"{cycle_index}.{doc_index}"
    
    def get_project_folder(self, project_name: str):
        """获取项目文件夹"""
        if self.folders:
            return self.folders.get_project_folder(project_name)
        return None
    
    def get_documents_folder(self, project_name: str):
        """获取文档文件夹"""
        if self.folders:
            return self.folders.get_documents_folder(project_name)
        return None
    
    def get_document_list_path(self, project_name: str):
        """获取文档清单路径"""
        if self.folders:
            return self.folders.get_document_list_path(project_name)
        return None
    
    def create_document_list(self, project_name: str, project_info: Dict,
                           requirements: Optional[Dict] = None) -> Dict:
        """创建文档清单"""
        if self.doc_lists:
            return self.doc_lists.create(project_name, project_info, requirements)
        return {'status': 'error', 'message': '文档清单模块不可用'}
    
    def export_requirements_to_json(self, project_config: Dict) -> str:
        """导出需求为JSON"""
        if self.requirements:
            return self.requirements.export_to_json(project_config)
        return '{}'
    
    def get_operation_logs(self, limit: int = 100, operation_type: str = None,
                          project: str = None) -> List[Dict]:
        """获取操作日志"""
        if self.logger:
            return self.logger.get_logs(limit, operation_type, project)
        return []
    
    # ==================== 配置版本管理 ====================
    
    def save_version(self, project_id: str, version_name: str, description: str = '') -> Dict:
        """保存当前配置为新版本"""
        if self.projects:
            return self.projects.save_version(project_id, version_name, description)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def list_versions(self, project_id: str) -> Dict:
        """列出所有配置版本"""
        if self.projects:
            return self.projects.list_versions(project_id)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def load_version(self, project_id: str, version_filename: str) -> Dict:
        """加载指定版本的配置"""
        if self.projects:
            return self.projects.load_version(project_id, version_filename)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def switch_version(self, project_id: str, version_filename: str) -> Dict:
        """切换到指定版本"""
        if self.projects:
            return self.projects.switch_version(project_id, version_filename)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def delete_version(self, project_id: str, version_filename: str) -> Dict:
        """删除指定版本"""
        if self.projects:
            return self.projects.delete_version(project_id, version_filename)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def export_version(self, project_id: str, version_filename: str) -> Dict:
        """导出指定版本的配置"""
        if self.projects:
            return self.projects.export_version(project_id, version_filename)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    # ==================== 需求模板管理 ====================
    
    def save_template(self, template_name: str, template_data: Dict, description: str = '') -> Dict:
        """保存需求模板"""
        if self.projects:
            return self.projects.save_template(template_name, template_data, description)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def list_templates(self) -> Dict:
        """列出所有需求模板"""
        if self.projects:
            return self.projects.list_templates()
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def load_template(self, template_id: str) -> Dict:
        """加载指定模板"""
        if self.projects:
            return self.projects.load_template(template_id)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def delete_template(self, template_id: str) -> Dict:
        """删除指定模板"""
        if self.projects:
            return self.projects.delete_template(template_id)
        return {'status': 'error', 'message': '项目管理模块不可用'}
    
    def get_documents(self, cycle: str = None, doc_name: str = None, project_id: str = None) -> List[Dict]:
        """获取文档列表
        
        Args:
            cycle: 周期名称
            doc_name: 文档名称
            project_id: 项目ID
            
        Returns:
            List[Dict]: 文档列表
        """
        result = []
        
        # 首先从项目配置中加载文档，确保数据最新
        if project_id:
            project_result = self.load_project(project_id)
            if project_result.get('status') == 'success':
                project_config = project_result.get('project')
                if project_config and 'documents' in project_config:
                    documents = project_config['documents']
                    # 遍历所有周期
                    for doc_cycle, cycle_info in documents.items():
                        # 过滤周期
                        if cycle and doc_cycle != cycle:
                            continue
                        # 检查是否有已上传的文档
                        if 'uploaded_docs' in cycle_info:
                            for doc in cycle_info['uploaded_docs']:
                                # 过滤文档名称
                                if doc_name and doc.get('doc_name') != doc_name:
                                    continue
                                # 确保文档有 ID
                                doc_id = doc.get('doc_id') or f"{doc_cycle}_{doc.get('doc_name')}_{doc.get('upload_time', '').replace(':', '_').replace('-', '_')}"
                                # 添加到结果列表
                                result.append({
                                    'id': doc_id,
                                    **doc
                                })
        
        # 如果项目配置中没有文档，再从内存中的documents_db获取
        if not result:
            for doc_id, doc in self.documents_db.items():
                # 过滤周期
                if cycle and doc.get('cycle') != cycle:
                    continue
                # 过滤文档名称
                if doc_name and doc.get('doc_name') != doc_name:
                    continue
                # 确保文档有 ID
                doc_copy = doc.copy()
                doc_copy['id'] = doc_id
                result.append(doc_copy)
        
        return result
    
    def delete_document(self, doc_id: str) -> Dict:
        """删除文档
        
        Args:
            doc_id: 文档ID
            
        Returns:
            Dict: 删除结果
        """
        try:
            # 首先从内存中的documents_db删除
            if doc_id in self.documents_db:
                del self.documents_db[doc_id]
                deleted = True
            else:
                deleted = False
            
            # 然后从项目配置文件中删除（同时更新内存中的数据）
            if hasattr(self, 'projects') and self.projects:
                # 首先更新内存中的projects_db数据
                for project_id, project_data in self.projects.projects_db.items():
                    if 'documents' in project_data:
                        for cycle, cycle_info in project_data['documents'].items():
                            if 'uploaded_docs' in cycle_info:
                                # 过滤掉要删除的文档
                                original_length = len(cycle_info['uploaded_docs'])
                                cycle_info['uploaded_docs'] = [
                                    doc for doc in cycle_info['uploaded_docs']
                                    if doc.get('doc_id') != doc_id
                                ]
                                if len(cycle_info['uploaded_docs']) != original_length:
                                    deleted = True
                
                # 然后遍历所有项目文件（包括子目录）
                projects_dir = self.config.projects_base_folder
                # 检查根目录下的JSON文件
                for project_file in projects_dir.glob('*.json'):
                    try:
                        import json
                        with open(project_file, 'r', encoding='utf-8') as f:
                            project_data = json.load(f)
                        
                        if 'documents' in project_data:
                            for cycle, cycle_info in project_data['documents'].items():
                                if 'uploaded_docs' in cycle_info:
                                    # 过滤掉要删除的文档
                                    original_length = len(cycle_info['uploaded_docs'])
                                    cycle_info['uploaded_docs'] = [
                                        doc for doc in cycle_info['uploaded_docs']
                                        if doc.get('doc_id') != doc_id
                                    ]
                                    if len(cycle_info['uploaded_docs']) != original_length:
                                        deleted = True
                                        # 保存更新后的项目配置
                                        with open(project_file, 'w', encoding='utf-8') as f:
                                            json.dump(project_data, f, ensure_ascii=False, indent=2)
                    except Exception as e:
                        logger.warning(f"处理项目文件 {project_file} 时出错: {e}")
                
                # 检查子目录中的project_config.json文件
                for project_dir in projects_dir.iterdir():
                    if project_dir.is_dir():
                        project_config_file = project_dir / 'project_config.json'
                        if project_config_file.exists():
                            try:
                                import json
                                with open(project_config_file, 'r', encoding='utf-8') as f:
                                    project_data = json.load(f)
                                
                                if 'documents' in project_data:
                                    for cycle, cycle_info in project_data['documents'].items():
                                        if 'uploaded_docs' in cycle_info:
                                            # 过滤掉要删除的文档
                                            original_length = len(cycle_info['uploaded_docs'])
                                            cycle_info['uploaded_docs'] = [
                                                doc for doc in cycle_info['uploaded_docs']
                                                if doc.get('doc_id') != doc_id
                                            ]
                                            if len(cycle_info['uploaded_docs']) != original_length:
                                                deleted = True
                                                # 保存更新后的项目配置
                                                with open(project_config_file, 'w', encoding='utf-8') as f:
                                                    json.dump(project_data, f, ensure_ascii=False, indent=2)
                            except Exception as e:
                                logger.warning(f"处理项目配置文件 {project_config_file} 时出错: {e}")
            
            # 从ProjectDataManager中删除文档
            if hasattr(self, 'data_manager') and self.data_manager:
                # 遍历所有项目
                projects_dir = self.config.projects_base_folder
                for project_dir in projects_dir.iterdir():
                    if project_dir.is_dir():
                        project_name = project_dir.name
                        # 加载项目配置
                        config = self.data_manager.load_full_config(project_name)
                        if config and 'documents' in config:
                            # 遍历所有周期
                            for cycle, cycle_info in config['documents'].items():
                                if 'uploaded_docs' in cycle_info:
                                    # 过滤掉要删除的文档
                                    original_length = len(cycle_info['uploaded_docs'])
                                    cycle_info['uploaded_docs'] = [
                                        doc for doc in cycle_info['uploaded_docs']
                                        if doc.get('doc_id') != doc_id
                                    ]
                                    if len(cycle_info['uploaded_docs']) != original_length:
                                        deleted = True
                                        # 保存更新后的项目配置
                                        self.data_manager.save_full_config(project_name, config)
                                        # 从文档索引中移除文档
                                        self.data_manager.remove_document_from_index(project_name, doc_id)
            
            # 检查并删除文档清单中的已归档文档
            if hasattr(self, 'doc_lists') and self.doc_lists:
                # 遍历所有项目
                projects_dir = self.config.projects_base_folder
                for project_dir in projects_dir.iterdir():
                    if project_dir.is_dir():
                        project_name = project_dir.name
                        doc_list = self.doc_lists.load(project_name)
                        if doc_list:
                            # 遍历所有周期和文档
                            if 'documents' in doc_list:
                                for cycle_name, cycle_docs in doc_list['documents'].items():
                                    for doc_info in cycle_docs:
                                        if 'files' in doc_info:
                                            # 过滤掉要删除的文件
                                            original_length = len(doc_info['files'])
                                            doc_info['files'] = [
                                                file_info for file_info in doc_info['files']
                                                if file_info.get('doc_id') != doc_id
                                            ]
                                            if len(doc_info['files']) != original_length:
                                                deleted = True
                                                # 保存更新后的文档清单
                                                self.doc_lists.save(project_name, doc_list)
            
            if deleted:
                return {'status': 'success'}
            else:
                return {'status': 'error', 'message': '文档不存在'}
        except Exception as e:
            logger.error(f"删除文档失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def update_document(self, doc_id: str, data: Dict) -> Dict:
        """更新文档元数据
        
        Args:
            doc_id: 文档ID
            data: 更新数据
            
        Returns:
            Dict: 更新结果
        """
        try:
            updated = False
            
            # 确保has_seal和has_seal_marked字段同步
            if 'has_seal' in data:
                data['has_seal_marked'] = data['has_seal']
            elif 'has_seal_marked' in data:
                data['has_seal'] = data['has_seal_marked']
            
            # 确保盖章相关字段的一致性
            if 'has_seal' in data and data['has_seal']:
                data['no_seal'] = False
            if 'no_seal' in data and data['no_seal']:
                data['has_seal'] = False
                data['has_seal_marked'] = False
            
            # 更新内存中的documents_db
            if doc_id in self.documents_db:
                # 只更新指定的字段，不覆盖整个文档
                for key, value in data.items():
                    self.documents_db[doc_id][key] = value
                updated = True
            
            # 更新项目配置文件中的数据
            if hasattr(self, 'projects') and self.projects:
                # 首先更新内存中的projects_db数据
                for project_id, project_data in self.projects.projects_db.items():
                    if 'documents' in project_data:
                        for cycle, cycle_info in project_data['documents'].items():
                            if 'uploaded_docs' in cycle_info:
                                for doc in cycle_info['uploaded_docs']:
                                    if doc.get('doc_id') == doc_id or doc.get('id') == doc_id:
                                        # 只更新指定的字段，保留其他属性
                                        for key, value in data.items():
                                            doc[key] = value
                                        updated = True
                                        # 保存更新后的项目配置
                                        self.projects.save(project_id, project_data)
                
                # 然后遍历所有项目文件（包括子目录）
                projects_dir = self.config.projects_base_folder
                # 检查根目录下的JSON文件
                for project_file in projects_dir.glob('*.json'):
                    try:
                        import json
                        # 尝试使用UTF-8编码读取，如果失败则使用gbk
                        try:
                            with open(project_file, 'r', encoding='utf-8') as f:
                                project_data = json.load(f)
                        except UnicodeDecodeError:
                            with open(project_file, 'r', encoding='gbk') as f:
                                project_data = json.load(f)
                        
                        if 'documents' in project_data:
                            for cycle, cycle_info in project_data['documents'].items():
                                if 'uploaded_docs' in cycle_info:
                                    for doc in cycle_info['uploaded_docs']:
                                        if doc.get('doc_id') == doc_id or doc.get('id') == doc_id:
                                            # 只更新指定的字段，保留其他属性
                                            for key, value in data.items():
                                                doc[key] = value
                                            updated = True
                                            # 保存更新后的项目配置
                                            with open(project_file, 'w', encoding='utf-8') as f:
                                                json.dump(project_data, f, ensure_ascii=False, indent=2)
                    except Exception as e:
                        logger.warning(f"处理项目文件 {project_file} 时出错: {e}")
                
                # 检查子目录中的project_config.json文件
                for project_dir in projects_dir.iterdir():
                    if project_dir.is_dir():
                        project_config_file = project_dir / 'project_config.json'
                        if project_config_file.exists():
                            try:
                                import json
                                # 尝试使用UTF-8编码读取，如果失败则使用gbk
                                try:
                                    with open(project_config_file, 'r', encoding='utf-8') as f:
                                        project_data = json.load(f)
                                except UnicodeDecodeError:
                                    with open(project_config_file, 'r', encoding='gbk') as f:
                                        project_data = json.load(f)
                                
                                if 'documents' in project_data:
                                    for cycle, cycle_info in project_data['documents'].items():
                                        if 'uploaded_docs' in cycle_info:
                                            for doc in cycle_info['uploaded_docs']:
                                                if doc.get('doc_id') == doc_id or doc.get('id') == doc_id:
                                                    # 只更新指定的字段，保留其他属性
                                                    for key, value in data.items():
                                                        doc[key] = value
                                                    updated = True
                                                    # 保存更新后的项目配置
                                                    with open(project_config_file, 'w', encoding='utf-8') as f:
                                                        json.dump(project_data, f, ensure_ascii=False, indent=2)
                            except Exception as e:
                                logger.warning(f"处理项目配置文件 {project_config_file} 时出错: {e}")
            
            # 从ProjectDataManager中更新文档
            if hasattr(self, 'data_manager') and self.data_manager:
                # 遍历所有项目
                projects_dir = self.config.projects_base_folder
                for project_dir in projects_dir.iterdir():
                    if project_dir.is_dir():
                        project_name = project_dir.name
                        # 加载项目配置
                        config = self.data_manager.load_full_config(project_name)
                        if config and 'documents' in config:
                            # 遍历所有周期
                            for cycle, cycle_info in config['documents'].items():
                                if 'uploaded_docs' in cycle_info:
                                    for doc in cycle_info['uploaded_docs']:
                                        if doc.get('doc_id') == doc_id or doc.get('id') == doc_id:
                                            # 只更新指定的字段，保留其他属性
                                            for key, value in data.items():
                                                doc[key] = value
                                            updated = True
                            # 保存更新后的项目配置
                            self.data_manager.save_full_config(project_name, config)
                            
                            # 同时更新文档索引
                            doc_index = self.data_manager.load_documents_index(project_name)
                            if 'documents' in doc_index and doc_id in doc_index['documents']:
                                # 只更新指定的字段，保留其他属性
                                for key, value in data.items():
                                    doc_index['documents'][doc_id][key] = value
                                self.data_manager.save_documents_index(project_name, doc_index)
            
            if updated:
                return {'status': 'success'}
            else:
                return {'status': 'error', 'message': '文档不存在'}
        except Exception as e:
            logger.error(f"更新文档失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def get_document_preview(self, doc_id: str) -> Dict:
        """获取文档预览
        
        Args:
            doc_id: 文档ID
            
        Returns:
            Dict: 预览结果
        """
        try:
            # 首先从内存中的documents_db查找
            doc = None
            if doc_id in self.documents_db:
                doc = self.documents_db[doc_id]
            else:
                # 从项目配置中查找
                if hasattr(self, 'projects') and self.projects:
                    for project_id, project_data in self.projects.projects_db.items():
                        if 'documents' in project_data:
                            for cycle, cycle_info in project_data['documents'].items():
                                if 'uploaded_docs' in cycle_info:
                                    for d in cycle_info['uploaded_docs']:
                                        if d.get('doc_id') == doc_id:
                                            doc = d
                                            break
                                    if doc:
                                        break
                            if doc:
                                break
            
            if doc:
                file_path = doc.get('file_path')
                if file_path:
                    from pathlib import Path
                    # 处理相对路径
                    file_path_obj = Path(file_path)
                    if not file_path_obj.is_absolute():
                        # 相对路径，相对于项目的uploads目录
                        project_name = doc.get('project_name')
                        if not project_name and hasattr(self, 'current_project') and self.current_project:
                            project_name = self.current_project.get('name')
                        
                        if project_name:
                            project_uploads_dir = self.get_documents_folder(project_name)
                            file_path_obj = project_uploads_dir / file_path
                        else:
                            # 如果没有项目名称，尝试使用绝对路径
                            # 检查文件是否存在于uploads目录中
                            if hasattr(self, 'config') and hasattr(self.config, 'upload_folder'):
                                upload_folder = self.config.upload_folder
                            else:
                                upload_folder = Path('uploads')
                            file_path_obj = upload_folder / file_path
                    
                    if not file_path_obj.exists():
                        return {'status': 'error', 'message': '文件不存在'}
                    from src.services.preview_service import PreviewService
                    return PreviewService.get_document_preview(str(file_path_obj))
                else:
                    return {'status': 'error', 'message': '文件路径不存在'}
            else:
                return {'status': 'error', 'message': '文档不存在'}
        except Exception as e:
            logger.error(f"获取文档预览失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def check_document_compliance(self, doc_id: str, requirement: str) -> Dict:
        """检查文档是否符合要求
        
        Args:
            doc_id: 文档ID
            requirement: 要求描述
            
        Returns:
            Dict: 检查结果
        """
        try:
            if doc_id in self.documents_db:
                doc = self.documents_db[doc_id]
                issues = []
                
                # 检查签名
                if '签名' in requirement or '签字' in requirement:
                    if not doc.get('signer') and not doc.get('no_signature'):
                        issues.append('缺少签名')
                
                # 检查盖章
                if '盖章' in requirement or '章' in requirement:
                    if not doc.get('has_seal_marked') and not doc.get('party_a_seal') and not doc.get('party_b_seal') and not doc.get('no_seal'):
                        issues.append('缺少盖章')
                
                return {
                    'is_compliant': len(issues) == 0,
                    'issues': issues
                }
            else:
                return {'is_compliant': False, 'issues': ['文档不存在']}
        except Exception as e:
            logger.error(f"检查文档合规性失败: {e}")
            return {'is_compliant': False, 'issues': [str(e)]}
    
    def batch_update_documents(self, doc_ids: List[str], action: str) -> Dict:
        """批量更新文档
        
        Args:
            doc_ids: 文档ID列表
            action: 操作类型
            
        Returns:
            Dict: 更新结果
        """
        try:
            success_count = 0
            
            # 从内存中的documents_db更新
            for doc_id in doc_ids:
                if doc_id in self.documents_db:
                    if action == 'mark_seal':
                        self.documents_db[doc_id]['has_seal_marked'] = True
                        self.documents_db[doc_id]['has_seal'] = True
                        self.documents_db[doc_id]['no_seal'] = False
                    elif action == 'mark_no_seal':
                        self.documents_db[doc_id]['no_seal'] = True
                        self.documents_db[doc_id]['has_seal_marked'] = False
                        self.documents_db[doc_id]['has_seal'] = False
                    elif action == 'mark_no_signature':
                        self.documents_db[doc_id]['no_signature'] = True
                    success_count += 1
            
            # 从项目配置中更新
            if hasattr(self, 'projects') and self.projects:
                # 首先使用项目管理器保存（处理内存中的数据）
                for project_id, project_data in self.projects.projects_db.items():
                    if 'documents' in project_data:
                        for cycle, cycle_info in project_data['documents'].items():
                            if 'uploaded_docs' in cycle_info:
                                for doc in cycle_info['uploaded_docs']:
                                    doc_id = doc.get('doc_id') or doc.get('id')
                                    if doc_id in doc_ids:
                                        if action == 'mark_seal':
                                            doc['has_seal_marked'] = True
                                            doc['has_seal'] = True
                                            doc['no_seal'] = False
                                        elif action == 'mark_no_seal':
                                            doc['no_seal'] = True
                                            doc['has_seal_marked'] = False
                                            doc['has_seal'] = False
                                        elif action == 'mark_no_signature':
                                            doc['no_signature'] = True
                        # 保存更新后的项目配置
                        self.projects.save(project_id, project_data)
                
                # 然后直接检查文件系统中的项目配置文件（确保所有文件都被更新）
                projects_dir = self.config.projects_base_folder
                # 检查子目录中的project_config.json文件
                for project_dir in projects_dir.iterdir():
                    if project_dir.is_dir():
                        project_config_file = project_dir / 'project_config.json'
                        if project_config_file.exists():
                            try:
                                import json
                                with open(project_config_file, 'r', encoding='utf-8') as f:
                                    project_data = json.load(f)
                                
                                if 'documents' in project_data:
                                    for cycle, cycle_info in project_data['documents'].items():
                                        if 'uploaded_docs' in cycle_info:
                                            for doc in cycle_info['uploaded_docs']:
                                                doc_id = doc.get('doc_id') or doc.get('id')
                                                if doc_id in doc_ids:
                                                    if action == 'mark_seal':
                                                        doc['has_seal_marked'] = True
                                                        doc['has_seal'] = True
                                                        doc['no_seal'] = False
                                                    elif action == 'mark_no_seal':
                                                        doc['no_seal'] = True
                                                        doc['has_seal_marked'] = False
                                                        doc['has_seal'] = False
                                                    elif action == 'mark_no_signature':
                                                        doc['no_signature'] = True
                                
                                # 保存更新后的项目配置
                                with open(project_config_file, 'w', encoding='utf-8') as f:
                                    json.dump(project_data, f, ensure_ascii=False, indent=2)
                            except Exception as e:
                                logger.warning(f"处理项目配置文件 {project_config_file} 时出错: {e}")
            
            # 从ProjectDataManager中更新文档
            if hasattr(self, 'data_manager') and self.data_manager:
                # 遍历所有项目
                projects_dir = self.config.projects_base_folder
                for project_dir in projects_dir.iterdir():
                    if project_dir.is_dir():
                        project_name = project_dir.name
                        # 加载项目配置
                        config = self.data_manager.load_full_config(project_name)
                        if config and 'documents' in config:
                            # 遍历所有周期
                            for cycle, cycle_info in config['documents'].items():
                                if 'uploaded_docs' in cycle_info:
                                    for doc in cycle_info['uploaded_docs']:
                                        doc_id = doc.get('doc_id') or doc.get('id')
                                        if doc_id in doc_ids:
                                            if action == 'mark_seal':
                                                doc['has_seal_marked'] = True
                                                doc['has_seal'] = True
                                                doc['no_seal'] = False
                                            elif action == 'mark_no_seal':
                                                doc['no_seal'] = True
                                                doc['has_seal_marked'] = False
                                                doc['has_seal'] = False
                                            elif action == 'mark_no_signature':
                                                doc['no_signature'] = True
                            # 保存更新后的项目配置
                            self.data_manager.save_full_config(project_name, config)
                            
                            # 同时更新文档索引
                            doc_index = self.data_manager.load_documents_index(project_name)
                            if 'documents' in doc_index:
                                for doc_id in doc_ids:
                                    if doc_id in doc_index['documents']:
                                        if action == 'mark_seal':
                                            doc_index['documents'][doc_id]['has_seal_marked'] = True
                                            doc_index['documents'][doc_id]['has_seal'] = True
                                            doc_index['documents'][doc_id]['no_seal'] = False
                                        elif action == 'mark_no_seal':
                                            doc_index['documents'][doc_id]['no_seal'] = True
                                            doc_index['documents'][doc_id]['has_seal_marked'] = False
                                            doc_index['documents'][doc_id]['has_seal'] = False
                                        elif action == 'mark_no_signature':
                                            doc_index['documents'][doc_id]['no_signature'] = True
                                self.data_manager.save_documents_index(project_name, doc_index)
            
            return {
                'status': 'success',
                'success_count': success_count,
                'total_count': len(doc_ids)
            }
        except Exception as e:
            logger.error(f"批量更新文档失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def batch_delete_documents(self, doc_ids: List[str]) -> Dict:
        """批量删除文档
        
        Args:
            doc_ids: 文档ID列表
            
        Returns:
            Dict: 删除结果
        """
        try:
            success_count = 0
            
            # 从内存中的documents_db删除
            for doc_id in doc_ids:
                if doc_id in self.documents_db:
                    del self.documents_db[doc_id]
                    success_count += 1
            
            # 从项目配置中删除
            if hasattr(self, 'projects') and self.projects:
                # 首先使用项目管理器保存（处理内存中的数据）
                for project_id, project_data in self.projects.projects_db.items():
                    if 'documents' in project_data:
                        for cycle, cycle_info in project_data['documents'].items():
                            if 'uploaded_docs' in cycle_info:
                                # 过滤掉要删除的文档
                                original_length = len(cycle_info['uploaded_docs'])
                                cycle_info['uploaded_docs'] = [
                                    doc for doc in cycle_info['uploaded_docs']
                                    if doc.get('doc_id') not in doc_ids
                                ]
                                if len(cycle_info['uploaded_docs']) != original_length:
                                    # 保存更新后的项目配置
                                    self.projects.save(project_id, project_data)
                
                # 然后直接检查文件系统中的项目配置文件（确保所有文件都被更新）
                projects_dir = self.config.projects_base_folder
                # 检查子目录中的project_config.json文件
                for project_dir in projects_dir.iterdir():
                    if project_dir.is_dir():
                        project_config_file = project_dir / 'project_config.json'
                        if project_config_file.exists():
                            try:
                                import json
                                with open(project_config_file, 'r', encoding='utf-8') as f:
                                    project_data = json.load(f)
                                
                                if 'documents' in project_data:
                                    for cycle, cycle_info in project_data['documents'].items():
                                        if 'uploaded_docs' in cycle_info:
                                            # 过滤掉要删除的文档
                                            original_length = len(cycle_info['uploaded_docs'])
                                            cycle_info['uploaded_docs'] = [
                                                doc for doc in cycle_info['uploaded_docs']
                                                if doc.get('doc_id') not in doc_ids
                                            ]
                                            if len(cycle_info['uploaded_docs']) != original_length:
                                                # 保存更新后的项目配置
                                                with open(project_config_file, 'w', encoding='utf-8') as f:
                                                    json.dump(project_data, f, ensure_ascii=False, indent=2)
                            except Exception as e:
                                logger.warning(f"处理项目配置文件 {project_config_file} 时出错: {e}")
                
                # 从ProjectDataManager中删除文档
                if hasattr(self, 'data_manager') and self.data_manager:
                    # 遍历所有项目
                    projects_dir = self.config.projects_base_folder
                    for project_dir in projects_dir.iterdir():
                        if project_dir.is_dir():
                            project_name = project_dir.name
                            # 加载项目配置
                            config = self.data_manager.load_full_config(project_name)
                            if config and 'documents' in config:
                                # 遍历所有周期
                                for cycle, cycle_info in config['documents'].items():
                                    if 'uploaded_docs' in cycle_info:
                                        # 过滤掉要删除的文档
                                        original_length = len(cycle_info['uploaded_docs'])
                                        cycle_info['uploaded_docs'] = [
                                            doc for doc in cycle_info['uploaded_docs']
                                            if doc.get('doc_id') not in doc_ids
                                        ]
                                        if len(cycle_info['uploaded_docs']) != original_length:
                                            # 保存更新后的项目配置
                                            self.data_manager.save_full_config(project_name, config)
                                            # 从文档索引中移除文档
                                            for doc_id in doc_ids:
                                                self.data_manager.remove_document_from_index(project_name, doc_id)
            
            return {
                'status': 'success',
                'success_count': success_count,
                'total_count': len(doc_ids)
            }
        except Exception as e:
            logger.error(f"批量删除文档失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def get_categories(self, cycle: str, doc_name: str, project_id: str = None) -> List[str]:
        """获取分类列表
        
        Args:
            cycle: 周期名称
            doc_name: 文档名称
            project_id: 项目ID（可选，如果未提供则使用 current_project）
            
        Returns:
            List[str]: 分类列表
        """
        try:
            # 首先尝试使用新的数据管理器
            if hasattr(self, 'data_manager') and self.data_manager:
                # 获取项目名称
                project_name = None
                if project_id and hasattr(self, 'projects') and self.projects:
                    project_info = self.projects.get_by_id(project_id)
                    if project_info:
                        project_name = project_info.get('name')
                elif hasattr(self, 'current_project') and self.current_project:
                    project_name = self.current_project.get('name')
                
                if project_name:
                    return self.data_manager.get_categories_for_doc(project_name, cycle, doc_name)
            
            # 回退到旧方法
            project_config = None
            if project_id and hasattr(self, 'projects') and self.projects:
                project_config = self.projects.load(project_id)
            elif hasattr(self, 'current_project') and self.current_project:
                project_config = self.current_project
            
            if project_config:
                if 'documents' in project_config:
                    if cycle in project_config['documents']:
                        if 'categories' in project_config['documents'][cycle]:
                            if doc_name in project_config['documents'][cycle]['categories']:
                                return project_config['documents'][cycle]['categories'][doc_name]
            return []
        except Exception as e:
            logger.error(f"获取分类列表失败: {e}")
            return []
    
    def create_category(self, cycle: str, doc_name: str, category: str, project_id: str = None) -> Dict:
        """创建分类
        
        Args:
            cycle: 周期名称
            doc_name: 文档名称
            category: 分类名称
            project_id: 项目ID（可选，如果未提供则使用 current_project）
            
        Returns:
            Dict: 创建结果
        """
        try:
            # 首先尝试使用新的数据管理器
            if hasattr(self, 'data_manager') and self.data_manager:
                # 获取项目名称
                project_name = None
                if project_id and hasattr(self, 'projects') and self.projects:
                    project_info = self.projects.get_by_id(project_id)
                    if project_info:
                        project_name = project_info.get('name')
                elif hasattr(self, 'current_project') and self.current_project:
                    project_name = self.current_project.get('name')
                
                if project_name:
                    success = self.data_manager.add_category_for_doc(project_name, cycle, doc_name, category)
                    if success:
                        return {'status': 'success'}
            
            # 回退到旧方法
            project_config = None
            if project_id and hasattr(self, 'projects') and self.projects:
                project_config = self.projects.load(project_id)
            elif hasattr(self, 'current_project') and self.current_project:
                project_config = self.current_project
                project_id = project_config.get('id')
            
            if not project_config:
                return {'status': 'error', 'message': '项目未加载'}
            
            if 'documents' not in project_config:
                project_config['documents'] = {}
            if cycle not in project_config['documents']:
                project_config['documents'][cycle] = {'categories': {}}
            if 'categories' not in project_config['documents'][cycle]:
                project_config['documents'][cycle]['categories'] = {}
            if doc_name not in project_config['documents'][cycle]['categories']:
                project_config['documents'][cycle]['categories'][doc_name] = []
            
            # 检查分类是否已存在
            if category not in project_config['documents'][cycle]['categories'][doc_name]:
                project_config['documents'][cycle]['categories'][doc_name].append(category)
                # 保存项目配置
                if hasattr(self, 'projects') and self.projects and project_id:
                    self.projects.save(project_id, project_config)
                # 更新 current_project
                if hasattr(self, 'current_project'):
                    self.current_project = project_config
            
            return {'status': 'success'}
        except Exception as e:
            logger.error(f"创建分类失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def delete_category(self, cycle: str, doc_name: str, category: str, project_id: str = None) -> Dict:
        """删除分类
        
        Args:
            cycle: 周期名称
            doc_name: 文档名称
            category: 分类名称
            project_id: 项目ID（可选，如果未提供则使用 current_project）
            
        Returns:
            Dict: 删除结果
        """
        try:
            # 首先尝试使用新的数据管理器
            if hasattr(self, 'data_manager') and self.data_manager:
                # 获取项目名称
                project_name = None
                if project_id and hasattr(self, 'projects') and self.projects:
                    project_info = self.projects.get_by_id(project_id)
                    if project_info:
                        project_name = project_info.get('name')
                elif hasattr(self, 'current_project') and self.current_project:
                    project_name = self.current_project.get('name')
                
                if project_name:
                    success = self.data_manager.remove_category_for_doc(project_name, cycle, doc_name, category)
                    if success:
                        return {'status': 'success'}
            
            # 回退到旧方法
            project_config = None
            if project_id and hasattr(self, 'projects') and self.projects:
                project_config = self.projects.load(project_id)
            elif hasattr(self, 'current_project') and self.current_project:
                project_config = self.current_project
                project_id = project_config.get('id')
            
            if not project_config:
                return {'status': 'error', 'message': '项目未加载'}
            
            if 'documents' in project_config:
                if cycle in project_config['documents']:
                    if 'categories' in project_config['documents'][cycle]:
                        if doc_name in project_config['documents'][cycle]['categories']:
                            categories = project_config['documents'][cycle]['categories'][doc_name]
                            if category in categories:
                                categories.remove(category)
                                # 保存项目配置
                                if hasattr(self, 'projects') and self.projects and project_id:
                                    self.projects.save(project_id, project_config)
                                # 更新 current_project
                                if hasattr(self, 'current_project'):
                                    self.current_project = project_config
            
            return {'status': 'success'}
        except Exception as e:
            logger.error(f"删除分类失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def replace_document(self, doc_id: str, file, data: Dict) -> Dict:
        """替换文档
        
        Args:
            doc_id: 文档ID
            file: 文件对象
            data: 更新数据
            
        Returns:
            Dict: 替换结果
        """
        try:
            if doc_id in self.documents_db:
                doc = self.documents_db[doc_id]
                cycle = doc.get('cycle')
                doc_name = doc.get('doc_name')
                
                if not cycle or not doc_name:
                    return {'status': 'error', 'message': '文档信息不完整'}
                
                # 上传新文件
                result = self.upload_document(file, cycle, doc_name, **data)
                if result['status'] == 'success':
                    # 删除旧文档
                    del self.documents_db[doc_id]
                    return {'status': 'success', 'message': '文档替换成功'}
                else:
                    return result
            else:
                return {'status': 'error', 'message': '文档不存在'}
        except Exception as e:
            logger.error(f"替换文档失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def extract_zipfile(self, zip_path: str, project_config: Dict) -> Dict:
        """解压ZIP文件并自动匹配文档
        
        Args:
            zip_path: ZIP文件路径
            project_config: 项目配置
            
        Returns:
            Dict: 解压和匹配结果
        """
        try:
            from .zip_matcher import create_matcher
            
            # 创建匹配器
            matcher = create_matcher({'upload_folder': 'uploads'})
            
            # 获取项目名称
            project_name = project_config.get('name') if project_config else None
            
            # 执行匹配
            result = matcher.extract_and_match(
                zip_path, 
                project_config,
                project_name=project_name
            )
            
            # 保存更新后的项目配置
            if project_config and result.get('status') == 'success':
                project_id = project_config.get('id')
                if project_id and self.projects:
                    self.projects.save(project_id, project_config)
            
            return result
        except Exception as e:
            logger.error(f"解压ZIP文件失败: {e}")
            return {'status': 'error', 'message': str(e)}

    def smart_recognize_document(self, file_path: str, party_a: str = '', party_b: str = '', 
                                  attributes_to_recognize: Dict = None) -> Dict:
        """智能识别文档属性（签章、盖章等）
        
        Args:
            file_path: 文件路径
            party_a: 甲方名称
            party_b: 乙方名称
            attributes_to_recognize: 需要识别的属性配置
            
        Returns:
            Dict: 识别结果
        """
        try:
            from pathlib import Path
            
            file_path_obj = Path(file_path)
            file_ext = file_path_obj.suffix.lower()
            
            # 默认识别所有属性
            if attributes_to_recognize is None:
                attributes_to_recognize = {
                    'doc_date': True,
                    'sign_date': True,
                    'signer': True,
                    'has_seal': True,
                    'party_a_seal': True,
                    'party_b_seal': True,
                    'no_seal': True,
                    'no_signature': True,
                    'other_seal': True,
                    'doc_number': True
                }
            
            result = {
                'doc_date': '',
                'sign_date': '',
                'signer': '',
                'has_seal': False,
                'party_a_seal': False,
                'party_b_seal': False,
                'no_seal': False,
                'no_signature': False,
                'other_seal': '',
                'doc_number': ''
            }
            
            # 根据文件类型选择不同的识别方法
            if file_ext == '.pdf':
                result = self._recognize_pdf(file_path, party_a, party_b, attributes_to_recognize)
            elif file_ext in ['.doc', '.docx']:
                result = self._recognize_word(file_path, party_a, party_b, attributes_to_recognize)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                result = self._recognize_image(file_path, party_a, party_b, attributes_to_recognize)
            else:
                # 对于其他格式，尝试通用文本识别
                result = self._recognize_text(file_path, party_a, party_b, attributes_to_recognize)
            
            return {
                'status': 'success',
                'data': result
            }
        except Exception as e:
            logger.error(f"智能识别文档失败: {e}")
            return {
                'status': 'error',
                'message': str(e),
                'data': {
                    'doc_date': '',
                    'sign_date': '',
                    'signer': '',
                    'has_seal': False,
                    'party_a_seal': False,
                    'party_b_seal': False,
                    'no_seal': False,
                    'no_signature': False,
                    'other_seal': '',
                    'doc_number': ''
                }
            }
    
    def _recognize_pdf(self, file_path: str, party_a: str, party_b: str, 
                       attributes_to_recognize: Dict = None) -> Dict:
        """识别PDF文档"""
        result = {
            'doc_date': '',
            'sign_date': '',
            'signer': '',
            'has_seal': False,
            'party_a_seal': False,
            'party_b_seal': False,
            'no_seal': False,
            'no_signature': False,
            'other_seal': '',
            'doc_number': ''
        }
        
        # 默认全部识别
        if attributes_to_recognize is None:
            attributes_to_recognize = {k: True for k in result.keys()}
        
        try:
            import PyPDF2
            from datetime import datetime
            import re
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ''
                
                # 提取所有页面的文本
                for page in pdf_reader.pages:
                    text += page.extract_text() + '\n'
                
                # 识别日期（多种格式）
                if attributes_to_recognize.get('doc_date') or attributes_to_recognize.get('sign_date'):
                    date_patterns = [
                        r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})[日]?',
                        r'(\d{4})(\d{2})(\d{2})',
                        r'(\d{2})[年/-](\d{1,2})[月/-](\d{1,2})[日]?'
                    ]
                    
                    dates = []
                    for pattern in date_patterns:
                        matches = re.findall(pattern, text)
                        for match in matches:
                            if len(match) == 3:
                                year, month, day = match
                                if len(year) == 2:
                                    year = '20' + year
                                try:
                                    date_str = f"{year}-{int(month):02d}-{int(day):02d}"
                                    dates.append(date_str)
                                except:
                                    pass
                    
                    if dates and attributes_to_recognize.get('doc_date'):
                        result['doc_date'] = dates[0]
                    if len(dates) > 1 and attributes_to_recognize.get('sign_date'):
                        result['sign_date'] = dates[1]
                
                # 识别签字人（常见关键词附近的人名）
                if attributes_to_recognize.get('signer'):
                    signer_patterns = [
                        r'(?:签字|签名|签署)[：:]\s*([^\n]{1,20})',
                        r'(?:负责人|经办人|审核人)[：:]\s*([^\n]{1,20})',
                        r'(?:甲方代表|乙方代表)[：:]\s*([^\n]{1,20})'
                    ]
                    
                    for pattern in signer_patterns:
                        match = re.search(pattern, text)
                        if match:
                            signer = match.group(1).strip()
                            # 过滤掉过长的匹配
                            if len(signer) < 20 and not any(char in signer for char in ['【', '】', '[', ']']):
                                result['signer'] = signer
                                break
                
                # 识别盖章信息
                if attributes_to_recognize.get('has_seal') or attributes_to_recognize.get('party_a_seal') or attributes_to_recognize.get('party_b_seal'):
                    if party_a and party_a in text:
                        result['has_seal'] = True
                        if attributes_to_recognize.get('party_a_seal'):
                            result['party_a_seal'] = True
                
                if party_b and party_b in text:
                    result['has_seal'] = True
                    result['party_b_seal'] = True
                
                # 检查是否有"章"、"印章"等关键词
                seal_keywords = ['盖章', '印章', '公章', '专用章', '合同章']
                for keyword in seal_keywords:
                    if keyword in text:
                        result['has_seal'] = True
                        break
                
                # 检查是否涉及签字/盖章
                if '不涉及签字' in text or '无需签字' in text:
                    result['no_signature'] = True
                
                if '不涉及盖章' in text or '无需盖章' in text:
                    result['no_seal'] = True
                
                # 识别其他盖章信息
                other_seal_patterns = [
                    r'(监理单位|项目章|部门章|财务章)[：:]\s*([^\n]{1,30})',
                    r'(见证单位|检测单位|设计单位)[：:]\s*([^\n]{1,30})'
                ]
                
                other_seals = []
                for pattern in other_seal_patterns:
                    matches = re.findall(pattern, text)
                    for match in matches:
                        other_seals.append(f"{match[0]}: {match[1].strip()}")
                
                if other_seals:
                    result['other_seal'] = '; '.join(other_seals[:3])  # 最多3个
                
        except Exception as e:
            logger.warning(f"PDF识别失败: {e}")
        
        return result
    
    def _recognize_word(self, file_path: str, party_a: str, party_b: str) -> Dict:
        """识别Word文档"""
        result = {
            'doc_date': '',
            'sign_date': '',
            'signer': '',
            'has_seal': False,
            'party_a_seal': False,
            'party_b_seal': False,
            'no_seal': False,
            'no_signature': False,
            'other_seal': ''
        }
        
        try:
            from docx import Document
            import re
            from datetime import datetime
            
            doc = Document(file_path)
            text = ''
            
            # 提取所有段落的文本
            for para in doc.paragraphs:
                text += para.text + '\n'
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + '\n'
            
            # 使用与PDF相同的识别逻辑
            date_patterns = [
                r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})[日]?',
                r'(\d{4})(\d{2})(\d{2})',
                r'(\d{2})[年/-](\d{1,2})[月/-](\d{1,2})[日]?'
            ]
            
            dates = []
            for pattern in date_patterns:
                matches = re.findall(pattern, text)
                for match in matches:
                    if len(match) == 3:
                        year, month, day = match
                        if len(year) == 2:
                            year = '20' + year
                        try:
                            date_str = f"{year}-{int(month):02d}-{int(day):02d}"
                            dates.append(date_str)
                        except:
                            pass
            
            if dates:
                result['doc_date'] = dates[0]
                if len(dates) > 1:
                    result['sign_date'] = dates[1]
            
            # 识别签字人
            signer_patterns = [
                r'(?:签字|签名|签署)[：:]\s*([^\n]{1,20})',
                r'(?:负责人|经办人|审核人)[：:]\s*([^\n]{1,20})',
                r'(?:甲方代表|乙方代表)[：:]\s*([^\n]{1,20})'
            ]
            
            for pattern in signer_patterns:
                match = re.search(pattern, text)
                if match:
                    signer = match.group(1).strip()
                    if len(signer) < 20 and not any(char in signer for char in ['【', '】', '[', ']']):
                        result['signer'] = signer
                        break
            
            # 识别盖章信息
            if party_a and party_a in text:
                result['has_seal'] = True
                result['party_a_seal'] = True
            
            if party_b and party_b in text:
                result['has_seal'] = True
                result['party_b_seal'] = True
            
            seal_keywords = ['盖章', '印章', '公章', '专用章', '合同章']
            for keyword in seal_keywords:
                if keyword in text:
                    result['has_seal'] = True
                    break
            
            if '不涉及签字' in text or '无需签字' in text:
                result['no_signature'] = True
            
            if '不涉及盖章' in text or '无需盖章' in text:
                result['no_seal'] = True
            
        except Exception as e:
            logger.warning(f"Word识别失败: {e}")
        
        return result
    
    def _recognize_image(self, file_path: str, party_a: str, party_b: str) -> Dict:
        """识别图片文档（使用OCR）"""
        result = {
            'doc_date': '',
            'sign_date': '',
            'signer': '',
            'has_seal': False,
            'party_a_seal': False,
            'party_b_seal': False,
            'no_seal': False,
            'no_signature': False,
            'other_seal': ''
        }
        
        try:
            # 尝试使用OCR（如果安装了相关库）
            try:
                from PIL import Image
                import pytesseract
                import re
                
                # 打开图片
                image = Image.open(file_path)
                
                # 使用OCR提取文本
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                
                # 使用与PDF相同的识别逻辑
                date_patterns = [
                    r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})[日]?',
                    r'(\d{4})(\d{2})(\d{2})'
                ]
                
                dates = []
                for pattern in date_patterns:
                    matches = re.findall(pattern, text)
                    for match in matches:
                        if len(match) == 3:
                            year, month, day = match
                            try:
                                date_str = f"{year}-{int(month):02d}-{int(day):02d}"
                                dates.append(date_str)
                            except:
                                pass
                
                if dates:
                    result['doc_date'] = dates[0]
                    if len(dates) > 1:
                        result['sign_date'] = dates[1]
                
                # 识别盖章信息
                if party_a and party_a in text:
                    result['has_seal'] = True
                    result['party_a_seal'] = True
                
                if party_b and party_b in text:
                    result['has_seal'] = True
                    result['party_b_seal'] = True
                
                seal_keywords = ['盖章', '印章', '公章', '专用章']
                for keyword in seal_keywords:
                    if keyword in text:
                        result['has_seal'] = True
                        break
                
            except ImportError:
                logger.warning("OCR库未安装，跳过图片识别")
                
        except Exception as e:
            logger.warning(f"图片识别失败: {e}")
        
        return result
    
    def _recognize_text(self, file_path: str, party_a: str, party_b: str) -> Dict:
        """通用文本识别"""
        result = {
            'doc_date': '',
            'sign_date': '',
            'signer': '',
            'has_seal': False,
            'party_a_seal': False,
            'party_b_seal': False,
            'no_seal': False,
            'no_signature': False,
            'other_seal': ''
        }
        
        try:
            # 尝试读取文件内容
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            import re
            
            # 简单的日期识别
            date_patterns = [
                r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})'
            ]
            
            dates = []
            for pattern in date_patterns:
                matches = re.findall(pattern, text)
                for match in matches:
                    if len(match) == 3:
                        year, month, day = match
                        try:
                            date_str = f"{year}-{int(month):02d}-{int(day):02d}"
                            dates.append(date_str)
                        except:
                            pass
            
            if dates:
                result['doc_date'] = dates[0]
            
            # 检查盖章信息
            if party_a and party_a in text:
                result['party_a_seal'] = True
            
            if party_b and party_b in text:
                result['party_b_seal'] = True
                
        except Exception as e:
            logger.warning(f"文本识别失败: {e}")
        
        return result
    
    def generate_report(self, project_config: Dict) -> Dict:
        """生成报告
        
        Args:
            project_config: 项目配置
            
        Returns:
            Dict: 报告数据
        """
        try:
            from datetime import datetime
            
            report = {
                'project_name': project_config.get('name', '未知项目'),
                'report_date': datetime.now().isoformat(),
                'cycles': [],
                'summary': {
                    'total_docs': 0,
                    'completed_docs': 0,
                    'pending_docs': 0
                }
            }
            
            # 统计信息
            total_docs = 0
            completed_docs = 0
            
            # 处理每个周期
            cycles = project_config.get('cycles', [])
            documents = project_config.get('documents', {})
            
            for cycle in cycles:
                cycle_info = documents.get(cycle, {})
                required_docs = cycle_info.get('required_docs', [])
                uploaded_docs = cycle_info.get('uploaded_docs', [])
                
                # 创建文档映射
                doc_map = {}
                for doc in uploaded_docs:
                    doc_name = doc.get('doc_name')
                    if doc_name not in doc_map:
                        doc_map[doc_name] = []
                    doc_map[doc_name].append(doc)
                
                cycle_report = {
                    'name': cycle,
                    'total_docs': len(required_docs),
                    'completed_docs': 0,
                    'documents': []
                }
                
                # 处理每个文档
                for req_doc in required_docs:
                    doc_name = req_doc.get('name')
                    requirement = req_doc.get('requirement', '')
                    
                    # 检查是否有上传的文档
                    uploaded = doc_map.get(doc_name, [])
                    has_uploaded = len(uploaded) > 0
                    
                    # 检查签名和盖章状态
                    has_signature = any(doc.get('signer') or doc.get('no_signature') for doc in uploaded)
                    has_seal = any(doc.get('has_seal') or doc.get('has_seal_marked') or doc.get('party_a_seal') or doc.get('party_b_seal') or doc.get('no_seal') for doc in uploaded)
                    
                    # 检查是否满足要求
                    is_completed = has_uploaded
                    if requirement:
                        if '签名' in requirement and not has_signature:
                            is_completed = False
                        if '盖章' in requirement and not has_seal:
                            is_completed = False
                    
                    # 收集文档的备注信息
                    notes = []
                    for doc in uploaded:
                        if 'notes' in doc and doc['notes']:
                            notes.append(doc['notes'])
                        # 检查其他可能存储备注的字段
                        if 'note' in doc and doc['note']:
                            notes.append(doc['note'])
                        if 'doc_note' in doc and doc['doc_note']:
                            notes.append(doc['doc_note'])
                        if 'comment' in doc and doc['comment']:
                            notes.append(doc['comment'])
                    
                    doc_report = {
                        'name': doc_name,
                        'requirement': requirement,
                        'status': 'completed' if is_completed else 'pending',
                        'has_uploaded': has_uploaded,
                        'has_signature': has_signature,
                        'has_seal': has_seal,
                        'uploaded_count': len(uploaded),
                        'notes': notes  # 添加备注信息
                    }
                    
                    cycle_report['documents'].append(doc_report)
                    if is_completed:
                        cycle_report['completed_docs'] += 1
                        completed_docs += 1
                    total_docs += 1
                
                report['cycles'].append(cycle_report)
            
            # 更新汇总信息
            report['summary']['total_docs'] = total_docs
            report['summary']['completed_docs'] = completed_docs
            report['summary']['pending_docs'] = total_docs - completed_docs
            
            return report
            
        except Exception as e:
            logger.error(f"生成报告失败: {e}")
            return {
                'project_name': '未知项目',
                'report_date': datetime.now().isoformat(),
                'cycles': [],
                'summary': {
                    'total_docs': 0,
                    'completed_docs': 0,
                    'pending_docs': 0
                },
                'error': str(e)
            }


# 全局单例
_manager_instance: Optional[DocumentManager] = None


def get_manager(config: Optional[Dict] = None) -> DocumentManager:
    """获取文档管理器单例
    
    Args:
        config: 配置字典
        
    Returns:
        DocumentManager: 文档管理器实例
    """
    global _manager_instance
    if _manager_instance is None or config is not None:
        _manager_instance = DocumentManager(config)
    return _manager_instance


def reset_manager():
    """重置管理器单例"""
    global _manager_instance
    _manager_instance = None
